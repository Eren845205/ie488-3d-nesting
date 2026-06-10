"""make_docx.py — Sunumu telefon-dostu tek Word belgesine (SUNUM.docx) döker.

Her slayt: başlık + içerik maddeleri + görsel + sarı kutuda "SEN SÖYLE" notları.
Sonda: sık sorulabilecek sorular eki.

Çalıştır:  python scripts/make_docx.py
Çıktı   :  SUNUM.docx (proje kök dizini)
"""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_ROOT = Path(__file__).resolve().parent.parent
_RES = _ROOT / "results"

ACCENT = RGBColor(0x25, 0x63, 0xEB)
GREEN = RGBColor(0x05, 0x96, 0x59)
INK = RGBColor(0x1E, 0x29, 0x3B)


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _say_box(doc, text):
    """Sarı zeminli, 'SEN SÖYLE' etiketli not kutusu."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _shade(cell, "FEF9C3")
    p0 = cell.paragraphs[0]
    r0 = p0.add_run("🎤 SEN SÖYLE")
    r0.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = RGBColor(0x92, 0x6A, 0x00)
    for line in text.split("\n"):
        p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11); run.font.color.rgb = INK
        if line.startswith("?"):
            run.bold = True; run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
    doc.add_paragraph()


def _bullets(doc, items):
    for text, color, bold in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.size = Pt(12); r.font.color.rgb = color; r.bold = bold


def _img(doc, path, width):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _heading(doc, n, title):
    h = doc.add_heading(level=1)
    run = h.add_run(f"Slayt {n} — {title}")
    run.font.color.rgb = ACCENT


doc = Document()
# Sayfa kenar boşluklarını biraz daralt (telefon okuması için içerik genişlesin)
for sec in doc.sections:
    sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)

# --- Kapak ---
t = doc.add_heading(level=0)
tr = t.add_run("Eklemeli İmalatta 2B Nesting Optimizasyonu")
tr.font.color.rgb = INK
sub = doc.add_paragraph()
sr = sub.add_run("IE 488 Dönem Projesi · Hoca sunumu (telefon-dostu sürüm)\n"
                 "Baseline (BL + BFD) → Simulated Annealing · 220×220 mm FDM tablası")
sr.font.size = Pt(12); sr.italic = True; sr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
doc.add_paragraph()

# --- 1 ---
_heading(doc, 1, "Açılış")
_bullets(doc, [
    ("AM'de nesting: parçaları build tablasına en az boşlukla dizmek.", INK, False),
    ("Bir Endüstri Mühendisliği optimizasyon problemi.", INK, False),
])
_say_box(doc,
"1) AM'de nesting problemine baktım: parçaları tablaya en az boşlukla dizmek — bir Endüstri Mühendisliği optimizasyon problemi.\n"
"2) Önce iki temel yöntem yazdım (Bottom-Left + Best-Fit Decreasing), sonra üstüne bir Simulated Annealing optimizasyon katmanı geliştirdim.\n"
"3) Medium sette doluluğu %86.9'dan %90.9'a çıkardım.\n"
"NOT: 'en optimize' deme — 'doluluğu artıracak şekilde' de.")

# --- 2 ---
_heading(doc, 2, "Problem: Nesting")
_bullets(doc, [
    ("Bir 3D yazıcı tablasına birden çok parçayı yerleştir.", INK, False),
    ("Amaç: alan kullanımını maksimize et → atığı, malzemeyi, baskı süresini azalt.", INK, False),
    ("Klasik IE optimizasyon problemi: kısıtlı kaynağın (tabla) en verimli kullanımı.", INK, False),
    ("Kapsam: NfAM (yalnız yerleştirme) · 2B · tek tabla.", INK, False),
])
_say_box(doc,
"'Nesting = parçaları tablaya dizme. Amaç en az boşluk → en az atık ve süre. Kısıtlı kaynağı verimli kullanma, tam bir IE konusu.'\n"
"Benzetme: valize en az boşlukla eşya koymak / otoparka en çok araba sığdırmak.\n"
"?Neden önemli? Az atık = az malzeme + az baskı süresi + az maliyet.")

# --- 3 ---
_heading(doc, 3, "Kapsam ve Yaklaşım")
_bullets(doc, [
    ("Parça temsili: AABB — her parça bir dikdörtgen (en × boy). Rotation {0°, 90°}.", INK, False),
    ("Tabla: 220×220 mm FDM yatağı, parçalar arası 1 mm boşluk.", INK, False),
    ("Literatür: Tang vd. 2025 review'inin E\\P\\S kutusu.", INK, False),
    ("Araujo 2019 taksonomisi de bu sınıf için Bottom-Left tipi sezgiseli öneriyor.", INK, False),
])
_say_box(doc,
"'Parçayı dikdörtgenle temsil ettim (AABB) — en yaygın, en ucuz. 2B seçtim çünkü FDM'de izdüşüm yeterli, aynı iskelet 3D'ye taşınabilir.'\n"
"?E\\P\\S ne? E=dikdörtgen temsil (AABB), P=sadece yerleştir, S=tek algoritma.\n"
"?Neden AABB? En basit, en hızlı, 'en eski en yaygın'. Polygon sonraki adım.")

# --- 4 ---
_heading(doc, 4, "Adım 1 — Baseline (Constructive Sezgiseller)")
_bullets(doc, [
    ("Bottom-Left (BL): parçaları sol-alttan yerleştirir; en temel referans.", GREEN, True),
    ("Best-Fit Decreasing (BFD): en az artık bırakan yere koyar (skyline).", GREEN, True),
    ("İkisi de açgözlü: anlık iyiyi seçer, en iyiyi garanti etmez.", INK, False),
])
_img(doc, _RES / "bl_medium_on.png", 4.2)
cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = cap.add_run("Baseline (BL) — medium set, %86.9 · gri taralı alan = atık")
cr.font.size = Pt(9); cr.italic = True
_say_box(doc,
"'İki temel yöntem yazdım. Bottom-Left sol-alttan yerleştirir; Best-Fit Decreasing en az artık bırakan yere koyar. İkisi de açgözlü — en iyiyi garanti etmez. Bu sınırı aşmak için optimizasyon ekledim.'\n"
"?Greedy ne? Her adımda o an en iyiyi seçen, bütünü düşünmeyen yöntem.")

# --- 5 ---
_heading(doc, 5, "Adım 2 — Geliştirdiğim Optimizasyon (EN KRİTİK)")
_bullets(doc, [
    ("Gözlem: açgözlü kural, parçaların hangi sırada verildiğine duyarlı.", INK, False),
    ("Parça sırasını bir karar değişkeni yapıp en iyi sırayı arayabilirim.", GREEN, True),
    ("Yöntem: Simulated Annealing — içeride yine Bottom-Left'i kullanır.", INK, False),
])
_say_box(doc,
"'Fark ettim ki açgözlü kural parça sırasına duyarlı. O yüzden sırayı karar değişkeni yapıp en iyi sırayı aradım — Simulated Annealing ile.'\n"
"TEK-NEFESLİK TANIM (ezberle): 'Farklı parça sıralarını deneyip en çok sığdıranı arar; ara sıra kötü sırayı da kabul eder ki tek yere takılıp kalmasın.'\n"
"?Derin sorarsa (hamle/soğutma/kabul): 'Bu yönü geliştiriyorum; mantığı: sırayı değiştir, doluluğu ölç, daha iyiyi tut, ara sıra kötüyü de kabul et.' — bu kadar yeter, abartma.")

# --- 6 ---
_heading(doc, 6, "Sonuçlar — Doluluk Karşılaştırması")
_img(doc, _RES / "comparison_with_sa.png", 6.2)
_say_box(doc,
"'Yeşil çubuklar SA. Medium ve stress'te baseline'ı geçiyor; small'da eşit çünkü orada tüm parçalar zaten sığıyor — baseline optimal.'\n"
"x ekseni = parça setleri, y ekseni = doluluk %.")

# --- 7 ---
_heading(doc, 7, "Sonuçlar — Sayılarla")
rows = [
    ["Set", "Talep", "Baseline", "SA", "Yorum"],
    ["small", "%48", "47.6%", "47.6%", "Zaten optimal (hepsi sığıyor)"],
    ["medium", "%111", "86.9%", "90.9%", "Aramanın net kazancı: +4 puan"],
    ["stress", "%181", "82.7%", "91.7%", "%91.7'ye çıktı (çoğu sıralamadan)"],
]
tbl = doc.add_table(rows=4, cols=5); tbl.style = "Light Grid Accent 1"
for r in range(4):
    for c in range(5):
        cell = tbl.cell(r, c)
        run = cell.paragraphs[0].add_run(rows[r][c])
        run.font.size = Pt(10)
        if r == 0:
            run.bold = True
        if r == 2:
            run.bold = True
doc.add_paragraph()
_say_box(doc,
"'Talep = tüm parçaların alanı / tabla. %100'ü aşınca bir kısmı zaten sığmaz. small %48 hepsi sığar. medium %111 SA +4 puan (en temiz örnek). stress %181 → %91.7 ama dürüst: kazancın çoğu sıralamadan.'\n"
"?Doluluk %86 ama parça yerleşmedi? Doluluk=yerleşen alan/tabla; talep %111 olduğu için kısmı zaten sığmaz.\n"
"?+9 hep SA'dan mı? Hayır — stress ~+7 sıralamadan, ~+2 aramadan. Temiz örnek medium.")

# --- 8 ---
_heading(doc, 8, "SA Nasıl İyileştiriyor? (Yakınsama)")
_img(doc, _RES / "sa_convergence.png", 6.0)
_say_box(doc,
"'Bu grafik SA'nın çalıştığının kanıtı: iterasyonlar boyunca en iyi doluluk tırmanıyor. Medium baseline'ın altından başlayıp onu geçiyor — arama gerçekten daha iyi sıra buluyor. small düz çünkü zaten tavanda.'")

# --- 9 ---
_heading(doc, 9, "Sonraki Adımlar")
_bullets(doc, [
    ("(a) Genetik Algoritma — aynı mantıkla popülasyon tabanlı arama.", INK, False),
    ("(b) Polygon temsil + No-Fit-Polygon — gerçek kontur.", INK, False),
    ("(c) 3D voxel temsile geçiş — gerçek hacimsel nesting.", INK, False),
])
qp = doc.add_paragraph(); qr = qp.add_run("→ Hangisini önceliklendirmemi istersiniz?")
qr.bold = True; qr.font.size = Pt(13); qr.font.color.rgb = ACCENT
_say_box(doc,
"'Üç yöne gidebilirim.' (a/b/c tek cümle.) Sonra hocaya DÖN ve SOR: 'Hangisini önceliklendirmemi istersiniz?' — kapanış hamlesi, mutlaka sor.\n"
"?Neden şimdi yapmadın? 'Önce yönü sizinle netleştirmek istedim ki doğru şeye emek harcayayım.'")

# --- 10 ---
_heading(doc, 10, "Özet")
_bullets(doc, [
    ("Çalışan bir 2B AM nesting baseline'ı kurdum (BL + BFD, E\\P\\S).", INK, False),
    ("Üstüne bir Simulated Annealing optimizasyon katmanı geliştirdim.", GREEN, True),
    ("Kısıtlı sette doluluk: 86.9% → 90.9% (+4 puan, medium).", INK, True),
    ("Tekrar-üretilebilir (seed=42), 47 birim testiyle doğrulandı.", INK, False),
])
_say_box(doc, "Kapanış: baseline kurdum → üstüne SA geliştirdim → medium'da +4 puan, 47 testle doğrulandı. 'Teşekkürler, sorularınız?'")

# --- Ek: Sık sorulabilecek sorular ---
doc.add_page_break()
h = doc.add_heading(level=1); hr = h.add_run("EK — Hocanın Sorabileceği Sorular (hazır cevaplar)")
hr.font.color.rgb = ACCENT
qa = [
    ("Simulated Annealing nedir, nasıl çalışır?",
     "Farklı parça sıralarını deneyip en çok sığdıranı arayan yöntem; ara sıra kötü sırayı da kabul eder ki lokal optimuma takılmasın. Detay sorulursa: 'bu yönü geliştiriyorum.'"),
    ("Doluluk %86 ama parçalar yerleşmedi, çelişki değil mi?",
     "İki ayrı şey. Doluluk = yerleşen alan / tabla. Medium'da talep tablanın %111'i; bir kısmı zaten sığmaz. Optimizasyon boşlukları kapatıp daha çok parça sığdırdı."),
    ("Stress'te +9 puan hep SA'dan mı?",
     "Hayır. Stress'te kazancın ~+7'si sıralama kuralı değişiminden, ~+2'si SA aramasından. Aramanın asıl işi gördüğü temiz örnek medium (+4 puan)."),
    ("Baseline ile SA karşılaştırman adil mi?",
     "Adil — ikisi de aynı yerleştirme kuralını ve çakışma testini kullanıyor; tek fark SA'nın sırayı optimize etmesi."),
    ("Neden sadece dikdörtgen (AABB)?",
     "Review'de 'en eski, en yaygın' temsil, hesabı en ucuz. Gerçek kontur (polygon/NFP) sonraki adım."),
    ("Optimal'den ne kadar uzaksın?",
     "Şu an alt sınır yok; bir baseline+iyileştirme. Küçük setlerde ILP (exact) ile gap ölçmek net bir sonraki iş."),
    ("Sonuçlar tekrar üretilebilir mi?",
     "Evet, sabit tohumla (seed=42) aynı sonuç. 47 birim testi geçerli yerleşimi de doğruluyor (çakışma yok, sınır içinde)."),
    ("3D'ye nasıl taşınır?",
     "İskelet hazır: AABB (en,boy) → (en,boy,derinlik); yerleştirme height-map/voxel'e genişler; aynı SA mantığı 3D sırada da çalışır."),
    ("Rotation on/off neden tutarsız (bazen off daha iyi)?",
     "Greedy etkisi: rotation daha çok seçenek verir ama algoritma bunu kapsamlı arama yapmadan kullanır; sıra değişince sonuç düşebilir. Bu da neden optimizasyon eklediğimin gerekçesi."),
]
for q, a in qa:
    pq = doc.add_paragraph(); rq = pq.add_run("S: " + q); rq.bold = True; rq.font.color.rgb = INK
    pa = doc.add_paragraph(); ra = pa.add_run("C: " + a); ra.font.color.rgb = RGBColor(0x33, 0x44, 0x55)
    doc.add_paragraph()

out = _ROOT / "SUNUM.docx"
doc.save(str(out))
print(f"Yazildi: {out}")
