"""make_docx_3d.py — Aşama 2 (3B) sunumunu telefon-dostu Word belgesine döker.

Her slayt: başlık + içerik maddeleri + görsel + sarı kutuda "SEN SÖYLE" notları.
Sonda: sık sorulabilecek sorular eki (5_sunum_3d_konusma.md ile aynı içerik).

Çalıştır:  python scripts/make_docx_3d.py
Çıktı   :  SUNUM_3D.docx (proje kök dizini)
"""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_ROOT = Path(__file__).resolve().parent.parent
_RES = _ROOT / "results"

ACCENT = RGBColor(0x25, 0x63, 0xEB)
GREEN = RGBColor(0x05, 0x96, 0x59)
INK = RGBColor(0x1E, 0x29, 0x3B)
MUTED = RGBColor(0x64, 0x74, 0x8B)


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _say_box(doc, text):
    """Sarı zeminli, 'SEN SÖYLE' etiketli not kutusu."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _shade(cell, "FEF9C3")
    p0 = cell.paragraphs[0]
    r0 = p0.add_run("🎤 SEN SÖYLE")
    r0.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = RGBColor(0x92, 0x6A, 0x00)
    for line in text.split("\n"):
        p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11); run.font.color.rgb = INK
        if line.startswith("?"):
            run.bold = True; run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
    doc.add_paragraph()


def _bullets(doc, items):
    for text, color, bold in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.size = Pt(12); r.font.color.rgb = color; r.bold = bold


def _img(doc, path, width):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _caption(doc, text):
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(text)
    cr.font.size = Pt(9); cr.italic = True; cr.font.color.rgb = MUTED


def _heading(doc, n, title):
    h = doc.add_heading(level=1)
    run = h.add_run(f"Slayt {n} — {title}")
    run.font.color.rgb = ACCENT


doc = Document()
for sec in doc.sections:
    sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)

# --- Kapak ---
t = doc.add_heading(level=0)
tr = t.add_run("Eklemeli İmalatta 3B Voxel Nesting (Aşama 2)")
tr.font.color.rgb = INK
sub = doc.add_paragraph()
sr = sub.add_run("IE 488 Dönem Projesi · Hoca sunumu (telefon-dostu sürüm)\n"
                 "STL mesh → Voxelization → Order → DBLF + Simulated Annealing → .STL\n"
                 "220×220 mm taban · yükseklik minimize")
sr.font.size = Pt(12); sr.italic = True; sr.font.color.rgb = MUTED
doc.add_paragraph()

# --- 1 ---
_heading(doc, 1, "Açılış")
_bullets(doc, [
    ("2B iskelet (baseline + SA) hocanın direktifiyle 3B'ye taşındı.", INK, False),
    ("Çıktı: gerçek 3B modellerin yerleşimi + STL dosyası.", INK, False),
])
_say_box(doc,
"1) Geçen hafta 2B nesting'i göstermiştim: baseline sezgiseller + Simulated Annealing.\n"
"2) Yönlendirmenizle aynı iskeleti 3B'ye taşıdım: gerçek 3B modeller, voxelization, "
"yüksekliği minimize edilen kutu, çıktı STL.\n"
"3) İki senaryoda test ettim: direktifteki sette baseline çok güçlü çıktı; "
"kalabalık sette SA yüksekliği 3.4 mm düşürdü.\n"
"NOT: 'optimal çözdüm' deme — 'bu çözünürlükte optimuma yakın' de.")

# --- 2 ---
_heading(doc, 2, "Çıkış Noktası — Verilen Yön")
_img(doc, _ROOT / "WhatsApp Image 2026-06-09 at 21.28.07.jpeg", 4.6)
_caption(doc, "Toplantı notu (2026-06-09)")
_bullets(doc, [
    ("Pipeline: 3B mesh → Voxelization → Order → .STL", INK, False),
    ("Kutu: taban sabit, yükseklik serbest → yüksekliği minimize et.", GREEN, True),
    ("3 model, adetler ×10 / ×5 / ×15.", GREEN, True),
])
_say_box(doc,
"'Geçen toplantıda bu pipeline'ı çizmiştiniz: karmaşık bir 3B mesh alınacak, "
"voxel'leştirilecek, sıralanıp kutuya yerleştirilecek, sonuç STL olarak verilecek. "
"Kutunun tabanı sabit, yüksekliği serbest — amaç yüksekliği indirmek.'\n"
"Bu slayt 'direktifi doğru anladım mı' teyididir — hoca düzeltecekse burada düzeltir.")

# --- 3 ---
_heading(doc, 3, "Problem: 3B Strip Packing (Open-Dimension)")
_bullets(doc, [
    ("220×220 mm sabit taban; yükseklik sınırsız ama cezalı: amaç max yüksekliği "
     "minimize etmek.", INK, False),
    ("SLM/SLS toz yatağı mantığı: yükseklik ≈ baskı süresi + toz maliyeti.", INK, False),
    ("Literatür: Tang vd. 2025 N-IM (voxel) dalı — E\\P\\S'den N\\P\\S'ye yükseliş.", INK, False),
    ("İkincil metrik: packing density = parça hacmi / kullanılan zarf hacmi.", INK, False),
])
_say_box(doc,
"'Problem 3B strip packing: taban sabit, yüksekliği minimize ediyorum. Toz yataklı "
"sistemlerde yükseklik doğrudan baskı süresi ve toz maliyeti demek.'\n"
"?2B'den farkı? 2B'de doluluğu maksimize ediyordum; 3B'de tavan açık — amaç yüksekliği "
"indirmek. Aynı ailenin open-dimension üyesi.\n"
"?N-IM ne? Review'de 'gerçek geometri temsili' sınıfı (voxel/polygon); E-IM yalnız bbox idi.")

# --- 4 ---
_heading(doc, 4, "Parça Seti — 3 Model, 30 Parça")
_img(doc, _RES / "models_overview.png", 6.2)
_bullets(doc, [
    ("Sandalye ×10 (içbükey) · L-braket ×5 (asimetrik) · halka ×15 (delikli).", INK, False),
    ("Prosedürel üretim (scripts/generate_models.py) — pipeline model-bağımsız.", INK, False),
])
_say_box(doc,
"'Üç model seçtim: çizdiğiniz örneğe sadık kalarak bir sandalye, asimetrik bir L-braket "
"ve ortası delik bir halka. Adetler direktifteki gibi 10-5-15, toplam 30 parça.'\n"
"?Neden bu üçü? İçbükeylik, asimetri, delik — bbox ile temsil edilse büyük hacim israfı "
"olurdu; voxel temsilin değerini gösteriyorlar.\n"
"?Neden prosedürel? Tekrar üretilebilirlik + internete bağımlılık yok; herhangi bir STL de verilebilir.")

# --- 5 ---
_heading(doc, 5, "Adım 1 — Voxelization")
_bullets(doc, [
    ("Pitch 3.44 mm → taban 64×64 grid (sandalye ≈ 11×11×21 voxel).", INK, False),
    ("4 eksen-hizalı duruş: dik / yan × 0° / 90°.", INK, False),
    ("Her duruşun alt + üst kolon profilleri önceden çıkarılır.", INK, False),
    ("İç hacim dolduruluyor (fill) — kabuk grid metrikleri bozardı.", INK, False),
])
_say_box(doc,
"'Mesh'i 3.44 mm'lik voxel'lere bölüyorum; sandalye yaklaşık 11×11×21 voxel. Her parçanın "
"4 duruşu var. Yerleştirme hızlı olsun diye kolon profillerini önceden çıkarıyorum.'\n"
"?Neden 24 değil 4 oryantasyon? Parça stabilitesi + arama uzayını küçük tutmak; parametrik.\n"
"?Neden bu pitch? Hız/doğruluk dengesi — 64'lük grid dakika-altı demoya izin veriyor.")

# --- 6 ---
_heading(doc, 6, "Adım 2 — DBLF Baseline + Heightmap")
_bullets(doc, [
    ("Order: hacim-azalan sıralama (2B BFD kuralının 3B'si).", INK, False),
    ("Kutu durumu tek yükseklik haritası: H[x,y].", INK, False),
    ("Seçim: tüm (duruş, x, y) adaylarında en alçak TEPE — (z_tepe, z, y, x) "
     "sözlüksel minimumu = DBLF.", GREEN, True),
    ("Sınırlama (bilinçli): heightmap çıkıntı ALTINA parça sokamaz — hız için taviz.", INK, False),
])
_say_box(doc,
"'Kutuyu Tetris gibi tek yükseklik haritasıyla tutuyorum; parçanın oturacağı yer kolon "
"profillerinden tek formülle çıkıyor. Baseline kural: en alçak tepeyi veren yerleşimi seç.'\n"
"?Neden z değil z_tepe? Amaç tavanı indirmek — aynı z'de yatık duruş dik duruştan iyi.\n"
"?Sınırlama sorulursa: 'çıkıntı altı boşluk kullanılamıyor; bilinçli taviz. "
"Tam 3B collision sonraki aşamanın işi.'")

# --- 7 ---
_heading(doc, 7, "Adım 3 — Simulated Annealing (EN KRİTİK)")
_bullets(doc, [
    ("Karar değişkeni: sıra + her parçanın duruşu.", INK, False),
    ("Hamleler: swap / insert / reverse + duruş çevirme.", INK, False),
    ("Enerji = max yükseklik + 0.1 × RMS kolon yüksekliği.", GREEN, True),
    ("Başlangıç = DBLF çözümü → SA asla baseline'ın altına düşemez. Seed=42.", INK, False),
])
_say_box(doc,
"'SA'nın 2B iskeletini taşıdım; karar değişkeni artık sıra + duruş. "
"Enerji = max yükseklik + küçük bir kompaktlık terimi.'\n"
"?Derin soru gelirse RMS hikâyesi: 'İlk denemede kompaktlık terimi ortalama yükseklikti "
"ve SA hiç ilerlemiyordu. Sebep: dik duran halka ile yatan halka aynı kolon-hacmi kaplıyor "
"— ortalamaya göre özdeşler, gradyan yok. Kareli ortalamaya geçince yatık duruş kesin "
"avantajlı oldu ve arama çalıştı.'\n"
"?İkinci kalibrasyon: 't0=8 mm iken kabul oranı %86 — fiilen rastgele yürüyüş. "
"3 mm'ye indirince arama dengelendi.'")

# --- 8 ---
_heading(doc, 8, "Sonuç 1 — Direktif Seti (30 parça)")
_img(doc, _RES / "nesting3d_default.png", 5.6)
_caption(doc, "DBLF = SA: 44.7 mm — halkalar dik, sandalyeler yatık")
_bullets(doc, [
    ("Yükseklik 44.7 mm (13 voxel); SA 3000 iterasyonda dahi geçemedi.", INK, True),
    ("Greedy akıllıca: taban dolunca halkaları DİK dikiyor.", INK, False),
    ("Katı alt sınır 11 voxel — boşluk en çok 2 voxel.", INK, False),
    ("2B small set hikâyesinin 3B karşılığı: 'baseline doğrulandı'.", GREEN, True),
])
_say_box(doc,
"'Direktifteki sette baseline 44.7 mm buldu; SA'yı 3000 iterasyona kadar zorladım, "
"geçemedi. Bunu başarısızlık değil doğrulama olarak okuyorum.'\n"
"?Alt sınır sorulursa: 'Yatan sandalyenin en ince boyutu 11 voxel — hiçbir yerleşim "
"altına inemez. Ben 13'teyim; aradaki 2 voxel'in kapanabilirliği açık soru — dürüst boşluk bu.'")

# --- 9 ---
_heading(doc, 9, "Sonuç 2 — Stress Seti (42 parça): SA Kazancı")
_img(doc, _RES / "nesting3d_stress.png", 5.6)
_caption(doc, "SA sonrası stress yerleşimi: 55.0 mm")
_bullets(doc, [
    ("DBLF 58.4 mm → SA 55.0 mm (−3.4 mm).", GREEN, True),
    ("Density: 0.387 → 0.406.", GREEN, True),
    ("Adetler ×1.4 (14/7/21) — taban sıkışınca sıra + duruş kararları kritikleşiyor.", INK, False),
])
_say_box(doc,
"'Kalabalık versiyonda baseline 58.4'te kalıyor, SA 55.0'a indiriyor — bir voxel katmanı "
"kazandı, density de arttı.'\n"
"?Neden burada kazanıyor? Taban sıkışınca greedy'nin erken kararları sonrakileri kötü "
"yerlere itiyor; SA sırayı ve duruşları birlikte oynatıp tıkanıklığı çözüyor.")

# --- 10 ---
_heading(doc, 10, "Sonuçlar — DBLF vs SA Karşılaştırması")
_img(doc, _RES / "comparison_3d.png", 6.4)
_say_box(doc,
"'Sol grafik yükseklik (alçak = iyi): default'ta çubuklar eşit, stress'te SA 3.4 mm "
"indiriyor. Sağ grafik density (yüksek = iyi): stress'te 0.387'den 0.406'ya çıkıyor.'\n"
"2B sunumdaki karşılaştırma grafiğinin 3B karşılığı — yeşil SA, kalabalık sette geçiyor, "
"seyrek sette eşit çünkü baseline zaten tavanda.")

# --- 11 ---
_heading(doc, 11, "Sayılarla — Yakınsama ve Özet Tablo")
_img(doc, _RES / "sa3d_convergence_stress.png", 5.6)
rows = [
    ["Senaryo", "Parça", "DBLF", "SA", "Kazanç"],
    ["default", "30", "44.7 mm", "44.7 mm", "— (doğrulandı)"],
    ["stress", "42", "58.4 mm", "55.0 mm", "−3.4 mm · density +0.019"],
]
tbl = doc.add_table(rows=3, cols=5); tbl.style = "Light Grid Accent 1"
for r in range(3):
    for c in range(5):
        cell = tbl.cell(r, c)
        run = cell.paragraphs[0].add_run(rows[r][c])
        run.font.size = Pt(10)
        if r == 0 or r == 2:
            run.bold = True
doc.add_paragraph()
_say_box(doc,
"'Yakınsama grafiği SA'nın gerçekten arama yaptığının kanıtı: 90. iterasyon civarında "
"baseline'ın altına inen konfigürasyonu buluyor ve koruyor.'\n"
"?Neden tek basamak? Yükseklik voxel cinsinden — 3.44 mm'lik katmanlar. Pürüzsüz eğri "
"bu problemde olmaz; iyileşme katman katman gelir.")

# --- 12 ---
_heading(doc, 12, "Pipeline'ın Son Oku — .STL Çıktısı")
_bullets(doc, [
    ("Çözüm gerçek geometriye geri çevrilir: mesh + duruş + konum → tek STL sahnesi.", INK, False),
    ("results/nesting3d_result_default.stl · nesting3d_result_stress.stl", GREEN, True),
    ("Windows 3D Viewer / slicer ile açılır.", INK, False),
    ("Tek tık demo: DEMO_3D.bat (~1 dk).", INK, False),
    ("72 birim testi (47 eski 2B + 25 yeni 3B); 2B koduna dokunulmadı.", INK, False),
])
_say_box(doc,
"'Direktifteki son ok STL'di: çözümü voxel'de bırakmıyorum — gerçek mesh'lere dönüşüm "
"uygulayıp tek STL yazıyorum. Doğrudan slicer'a gidebilir.' İstenirse 3D Viewer'da aç.")

# --- 13 ---
_heading(doc, 13, "Sınırlamalar ve Sonraki Adımlar")
_bullets(doc, [
    ("Cavity doldurma: tam 3B voxel collision (sandalye altına halka).", INK, False),
    ("Daha ince pitch (96×96) — süre/kalite eğrisi.", INK, False),
    ("Lab'ın gerçek STL parçalarıyla test.", INK, False),
    ("Alt sınır analizi: 13 voxel optimal mi, 12 mümkün mü?", INK, False),
    ("GA karşılaştırması.", INK, False),
])
_say_box(doc,
"Beş yönü tek cümleyle say, sonra hocaya DÖN ve SOR: 'Hangisini önceliklendirmemi "
"istersiniz?' — kapanış hamlesi, mutlaka sor.\n"
"?Neden cavity ilk? En görünür sınırlama; çözülürse default sette de kazanç ihtimali doğar.")

# --- 14 ---
_heading(doc, 14, "Özet")
_bullets(doc, [
    ("Direktifteki pipeline uçtan uca çalışıyor: STL → voxel → order → kutu → STL.", GREEN, True),
    ("2B iskelet 3B'ye taşındı; kod tabanı ayrık, 2B'ye dokunulmadı.", INK, False),
    ("Direktif seti: 44.7 mm — baseline optimuma yakın (SA doğruladı).", INK, False),
    ("Stress: 58.4 → 55.0 mm — SA'nın net kazancı.", INK, True),
    ("Seed=42 · 72 test · tek tık demo + STL çıktıları.", INK, False),
])
_say_box(doc,
"Kapanış (4 cümle): 'Çizdiğiniz pipeline'ı uçtan uca kurdum. 2B iskelet 3B'ye taşındı. "
"Verdiğiniz sette baseline çok güçlü çıktı, SA doğruladı; kalabalık sette SA 3.4 mm "
"kazandırdı. Hepsi tekrar-üretilebilir, 72 testle doğrulanmış — STL çıktıları hazır.' "
"Teşekkür + soru al.")

# --- Ek: SSS ---
h = doc.add_heading(level=1)
run = h.add_run("EK — Sık Sorulabilecek Sorular")
run.font.color.rgb = ACCENT
faq = [
    ("Neden default sette SA kazanamadı?",
     "Baseline'ın çözümü gerçekten akıllı: taban dolunca halkaları dik dikiyor. 3000 "
     "iterasyonda geçilemedi. Katı alt sınır 11 voxel; ben 13'teyim — aradaki 2 voxel'in "
     "kapanabilirliği açık soru, dürüst boşluk bu."),
    ("Neden 4 oryantasyon, 24 değil?",
     "Parça stabilitesi + arama uzayını küçük tutmak; parametrik, artırılabilir."),
    ("Çıkıntı altına parça koyabiliyor musun?",
     "Hayır — heightmap'in bilinen sınırı, hız için bilinçli taviz. Tam 3B collision "
     "Aşama 3'ün ilk adayı."),
    ("Voxel kaybı ne kadar?",
     "Voxelization muhafazakâr: her eksende ~1 hücre taşma payı — güvenli yön, çakışma asla olmaz."),
    ("Süre?",
     "Baseline milisaniyeler; SA 1000 iterasyon ~25 sn; demo toplam ~1 dk."),
    ("Bu gerçek baskıya gider mi?",
     "STL çıktısı slicer'a doğrudan açılıyor; termal/support kısıtları modellenmedi — Aşama 3 konusu."),
]
for q, a in faq:
    p = doc.add_paragraph()
    rq = p.add_run("? " + q)
    rq.bold = True; rq.font.size = Pt(12); rq.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
    pa = doc.add_paragraph()
    ra = pa.add_run(a)
    ra.font.size = Pt(11); ra.font.color.rgb = INK

out = _ROOT / "SUNUM_3D.docx"
doc.save(str(out))
print(f"Yazildi: {out}")
